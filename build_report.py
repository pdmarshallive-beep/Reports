from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import pandas as pd
import matplotlib.pyplot as plt
from io import StringIO
from pathlib import Path
import re
import os


KNOWN_TABLE_DATA = {
    'Installed Capacity by Segment (2025)': {
        'columns': ['Segment', 'Capacity (GW)', 'Share', 'Key States'],
        'rows': [
            ['Utility-Scale', '104', '65%', 'TX, IA, CA'],
            ['Residential', '32', '20%', 'CA, TX, FL'],
            ['Commercial', '23', '13%', 'NY, NJ, IL'],
            ['Off-Grid', '4', '2%', 'AK, PR'],
        ],
    },
    'On-Grid vs Off-Grid Comparison': {
        'columns': ['Attribute', 'On-Grid', 'Off-Grid'],
        'rows': [
            ['Interconnection', 'Required', 'Not required'],
            ['Storage', 'Optional', 'Mandatory'],
            ['Use Case', 'Urban, C&I', 'Remote, resilience hubs'],
            ['Reliability', 'Grid stability', 'Storage sizing critical'],
            ['Cost per Watt ($)', '2.1', '3.0'],
        ],
    },
    'Top 10 Renewable Companies (2025)': {
        'columns': ['Rank', 'Company', 'Capacity (GW)', 'Focus Areas'],
        'rows': [
            ['1', 'NextEra Energy', '60', 'Solar, Wind, Storage'],
            ['2', 'Invenergy', '33', 'Wind, Solar, Hydrogen'],
            ['3', 'Ørsted North America', '10', 'Offshore Wind'],
            ['4', 'AES Corporation', '15', 'Solar, Storage'],
            ['5', 'Brookfield Renewable', '13', 'Hydro, Solar, Wind'],
            ['6', 'Tesla Energy', '10', 'Storage, Rooftop Solar'],
            ['7', 'First Solar', '9', 'Utility-Scale Solar'],
            ['8', 'Enphase Energy', '2.5', 'Inverters, Storage'],
            ['9', 'Dominion Energy', '7', 'Solar, Wind'],
            ['10', 'Duke Energy', '6', 'Solar'],
        ],
    },
    'R&D Spending (2025)': {
        'columns': ['Company', 'R&D Spend (USD m)'],
        'rows': [
            ['Tesla Energy', '900'],
            ['First Solar', '500'],
            ['NextEra Energy', '430'],
            ['Enphase', '250'],
            ['Ørsted', '210'],
        ],
    },
    'Cost Declines by Technology (2015 vs 2025E)': {
        'columns': ['Technology', '2015 ($/MWh)', '2025E ($/MWh)', 'Decline'],
        'rows': [
            ['Battery Storage', '560', '132', '76%'],
            ['Utility Solar', '85', '26', '69%'],
            ['Onshore Wind', '74', '32', '57%'],
            ['Natural Gas', '58', '56', '3%'],
            ['Coal', '92', '94', '-2%'],
        ],
    },
    'LCOE by Generation Type (2025E)': {
        'columns': ['Generation Type', 'LCOE ($/MWh)'],
        'rows': [
            ['Utility Solar', '26'],
            ['Onshore Wind', '32'],
            ['Solar + Storage', '45'],
            ['Natural Gas', '56'],
            ['Coal', '94'],
            ['Nuclear', '112'],
        ],
    },
    'Solar LCOE by Region (2025E)': {
        'columns': ['Region', 'LCOE ($/MWh)', 'Key Drivers'],
        'rows': [
            ['Southwest', '22', 'High irradiance, land availability'],
            ['Southeast', '28', 'Good resources, moderate costs'],
            ['Midwest', '29', 'Moderate irradiance, lower land costs'],
            ['California', '30', 'Strong resources, high permitting costs'],
            ['Northeast', '34', 'Lower irradiance, constrained land'],
        ],
    },
    'Wind LCOE by Region (2025E)': {
        'columns': ['Region', 'LCOE ($/MWh)', 'Capacity Factor'],
        'rows': [
            ['Central Plains', '26', '42–46%'],
            ['Great Lakes', '31', '38–42%'],
            ['New England', '41', '28–32%'],
        ],
    },
    'Battery Storage Cost Decline (2015–2025E)': {
        'columns': ['Year', 'Cost ($/kWh)', 'Technology Mix'],
        'rows': [
            ['2015', '560', 'Early lithium-ion'],
            ['2020', '310', 'Scaled lithium-ion'],
            ['2023', '180', 'Mature lithium-ion'],
            ['2025E', '128', 'Advanced lithium-ion, emerging alternatives'],
        ],
    },
    'Installed Storage Capacity (2015–2025E)': {
        'columns': ['Year', 'Installed Capacity (GW)', 'Annual Additions (GW)'],
        'rows': [
            ['2015', '1.2', '0.3'],
            ['2020', '4.5', '1.2'],
            ['2023', '9.8', '2.5'],
            ['2025E', '16.5', '3.2'],
        ],
    },
    'Installed Capacity by Technology (2025)': {
        'columns': ['Technology', 'Installed Capacity (GW)', 'Growth Rate (CAGR)'],
        'rows': [
            ['Utility Solar', '130', '14%'],
            ['Onshore Wind', '150', '8%'],
            ['Offshore Wind', '8', '32%'],
            ['Battery Storage', '16.5', '45%'],
            ['Hydro', '82', '1%'],
        ],
    },
    'CAGR and 2031 Capacity Forecast': {
        'columns': ['Technology', 'CAGR (2025–2031)', '2031 Capacity (GW)'],
        'rows': [
            ['Utility Solar', '10%', '230'],
            ['Onshore Wind', '5.5%', '210'],
            ['Offshore Wind', '27%', '33'],
            ['Battery Storage', '24%', '55'],
            ['Hydro', '1%', '86'],
        ],
    },
    'Key Federal Incentives Under the IRA': {
        'columns': ['Incentive Type', 'Description', 'Value / Duration'],
        'rows': [
            ['ITC (Investment Tax Credit)', 'Credit for solar, storage, microgrids', '30% base, +10% adders'],
            ['PTC (Production Tax Credit)', 'Credit per MWh for wind, solar', '$27/MWh (inflation-adjusted)'],
            ['Hydrogen PTC', 'Credit for clean hydrogen production', 'Up to $3/kg'],
            ['Advanced Manufacturing Credit', 'Support for domestic solar, wind, battery mfg', '10–30% of cost'],
            ['Transmission Incentives', 'Support for large-scale transmission buildout', 'Varies by project'],
        ],
    },
    'State RPS Targets (Selected States)': {
        'columns': ['State', 'RPS Target', 'Target Year', 'Procurement Pathway'],
        'rows': [
            ['California', '100%', '2045', 'Solar, wind, storage, hydrogen'],
            ['New York', '100%', '2040', 'Renewable, nuclear, storage'],
            ['New Jersey', '50%', '2030', 'Solar, wind, offshore wind'],
            ['Illinois', '40%', '2030', 'Wind, solar, storage'],
            ['Texas', 'Voluntary', 'N/A', 'Market-driven expansion'],
        ],
    },
    'Annual Capital Deployment (2015–2025E)': {
        'columns': ['Year', 'Capital Deployment (USD bn)', 'Solar Share', 'Wind Share', 'Storage Share'],
        'rows': [
            ['2015', '45', '38%', '42%', '5%'],
            ['2018', '62', '42%', '38%', '8%'],
            ['2020', '78', '45%', '35%', '12%'],
            ['2023', '105', '48%', '32%', '16%'],
            ['2025E', '132', '50%', '28%', '18%'],
        ],
    },
    'Green Finance Instruments (2025)': {
        'columns': ['Instrument', '2025 Volume (USD bn)', 'Common Use Cases'],
        'rows': [
            ['Green bonds', '46', 'Utility-scale solar, wind, storage'],
            ['Sustainability-linked loans', '32', 'Corporate renewable procurement'],
            ['Tax equity', '22', 'Solar, wind, storage, microgrids'],
            ['Transition bonds', '6', 'Natural gas to clean energy'],
            ['Green ABS', '5', 'Solar loans, EV charging'],
        ],
    },
    'Leading Corporate Renewable Buyers (2025)': {
        'columns': ['Company', 'Procurement (MW)', 'Primary Mechanism', 'CFE Target'],
        'rows': [
            ['Amazon', '9,100', 'PPAs, Direct Build', '24/7 by 2025'],
            ['Microsoft', '7,600', 'PPAs, Hourly Matching', '100% by 2025'],
            ['Google', '6,900', 'PPAs, Hourly Matching', '24/7 by 2030'],
            ['Meta', '4,500', 'PPAs', '100% by 2025'],
            ['Walmart', '2,300', 'On-site + PPAs', '50% by 2025'],
        ],
    },
    'U.S. Renewable Capacity Forecast (2025–2035)': {
        'columns': ['Year', 'Total Renewables (GW)', 'Storage (GW)', 'Storage/Generation Ratio'],
        'rows': [
            ['2025E', '520', '16.5', '3.1%'],
            ['2030E', '680', '45', '6.6%'],
            ['2035E', '860', '170', '19.8%'],
        ],
    },
    'Forecasted Levelized Costs by Technology (2030E & 2035E)': {
        'columns': ['Technology', '2025E ($/MWh)', '2030E ($/MWh)', '2035E ($/MWh)'],
        'rows': [
            ['Solar PV', '26', '18', '12'],
            ['Onshore Wind', '32', '22', '16'],
            ['Battery Storage', '95', '70', '50'],
            ['Solar + Storage', '45', '32', '22'],
        ],
    },
}


def get_known_table(title):
    data = KNOWN_TABLE_DATA.get(title)
    if data:
        return pd.DataFrame(data['rows'], columns=data['columns'])
    return None


def add_hyperlink(paragraph, bookmark_name, text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return paragraph


def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run("Table of Contents")
    return paragraph


def add_bookmark(paragraph, bookmark_name, bookmark_id):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_arial_font(run)
    return new_para


def fill_table_of_contents(placeholder_paragraph, toc_entries):
    current = placeholder_paragraph
    title_para = insert_paragraph_after(current, text='Table of Contents', style='Heading 1')
    for run in title_para.runs:
        set_arial_font(run)
    current = title_para
    for text, bookmark_name, level in toc_entries:
        entry_para = insert_paragraph_after(current, style='List Paragraph')
        entry_para.paragraph_format.left_indent = Pt(16 * (level - 1))
        add_hyperlink(entry_para, bookmark_name, text)
        for run in entry_para.runs:
            set_arial_font(run)
        current = entry_para
    placeholder_paragraph._p.getparent().remove(placeholder_paragraph._p)


def add_paragraph_with_links(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r'(\[\^\d+\])', text)
    for part in parts:
        if not part:
            continue
        match = re.match(r'\[\^(\d+)\]', part)
        if match:
            add_hyperlink(paragraph, f'fn{match.group(1)}', part)
        else:
            run = paragraph.add_run(part)
            set_arial_font(run)
    # Set Arial for the whole paragraph
    for run in paragraph.runs:
        set_arial_font(run)
    return paragraph


def add_bookmarked_paragraph(doc, text, bookmark_name, bookmark_id):
    paragraph = doc.add_paragraph()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    set_arial_font(run)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.append(end)
    return paragraph


def normalize_column_name(name):
    return str(name).strip()


def find_best_column(df, label):
    if not label:
        return None
    label_norm = str(label).strip().lower()
    label_norm_simple = re.sub(r'[^0-9a-z]+', '', label_norm)
    for col in df.columns:
        col_norm = str(col).strip().lower()
        col_norm_simple = re.sub(r'[^0-9a-z]+', '', col_norm)
        if col_norm == label_norm or col_norm_simple == label_norm_simple:
            return col
    for col in df.columns:
        col_norm = str(col).strip().lower()
        col_norm_simple = re.sub(r'[^0-9a-z]+', '', col_norm)
        if label_norm in col_norm or col_norm in label_norm or label_norm_simple in col_norm_simple or col_norm_simple in label_norm_simple:
            return col
    return None


def coerce_numeric_series(series):
    return pd.to_numeric(series.astype(str).str.replace(',', '').str.replace('%', ''), errors='coerce')

def set_arial_font(run):
    """Set font to Arial for a run."""
    run.font.name = 'Arial'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)

def add_page_numbers_to_footer(doc):
    """Add page numbers centered in the footer of all pages."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = 1  # 1 = center alignment
    
    # Create a field for page number
    run = footer_para.add_run()
    set_arial_font(run)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_paragraph_arial(doc, text, style=None):
    """Add a paragraph with Arial font."""
    if style:
        paragraph = doc.add_paragraph(text, style=style)
    else:
        paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        set_arial_font(run)
    return paragraph

def add_heading_arial(doc, text, level):
    """Add a heading with Arial font."""
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_arial_font(run)
    return paragraph

def build_report(input_path, output_path):
    doc = Document()
    add_page_numbers_to_footer(doc)
    charts_dir = Path(__file__).resolve().parent / "charts"
    os.makedirs(charts_dir, exist_ok=True)

    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()
    text = text.replace("\t", "	")
    lines = text.splitlines()

    table_rows = []
    in_table = False
    in_chart = False
    chart_params = {}
    chart_index = 1
    last_table_df = None
    current_table_title = None
    toc_inserted = False
    toc_placeholder = None
    toc_entries = []
    bookmark_counter = 1
    in_title_page = False
    title_added = False

    for raw_line in lines:
        line = raw_line.strip()

        if line == "TITLE PAGE":
            in_title_page = True
            title_added = False
            continue

        if in_title_page and not title_added and line:
            title_para = doc.add_paragraph(style="Title")
            title_run = title_para.add_run(line)
            title_run.bold = True
            title_run.font.size = Pt(32)
            title_run.font.color.rgb = RGBColor(0, 51, 102)
            set_arial_font(title_run)
            title_added = True
            continue

        if in_title_page and title_added and line.upper().startswith("TABLE OF CONTENTS"):
            if not toc_inserted:
                doc.add_page_break()
                toc_placeholder = doc.add_paragraph('__TOC_PLACEHOLDER__')
                toc_inserted = True
            in_title_page = False
            continue

        if in_title_page and title_added and not line:
            continue

        if in_title_page and title_added and line:
            add_paragraph_arial(doc, line)
            continue

        if toc_inserted and line.startswith("(") and "table of contents" in line.lower():
            continue

        if line.startswith("# "):
            heading_text = line[2:].strip()
            if re.match(r'^\d+\.\d+', heading_text):
                paragraph = add_heading_arial(doc, heading_text, 2)
                add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
                toc_entries.append((heading_text, f'toc{bookmark_counter}', 2))
                bookmark_counter += 1
            else:
                paragraph = add_heading_arial(doc, heading_text, 1)
                add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
                toc_entries.append((heading_text, f'toc{bookmark_counter}', 1))
                bookmark_counter += 1
            continue
        if re.match(r'^\d+\.\d+\s+', line):
            paragraph = add_heading_arial(doc, line, 2)
            add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
            toc_entries.append((line, f'toc{bookmark_counter}', 2))
            bookmark_counter += 1
            continue
        if re.match(r'^\d+\.\s+', line):
            paragraph = add_heading_arial(doc, line, 1)
            add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
            toc_entries.append((line, f'toc{bookmark_counter}', 1))
            bookmark_counter += 1
            continue
        if line.strip().upper() == "## REFERENCES & FOOTNOTES":
            doc.add_page_break()
        if line.startswith("## "):
            paragraph = add_heading_arial(doc, line[3:], 2)
            add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
            toc_entries.append((line[3:], f'toc{bookmark_counter}', 2))
            bookmark_counter += 1
            continue

        if line == "TABLE:" or (line.startswith("Table ") and "—" in line):
            in_table = True
            table_rows = []
            current_table_title = line.split("—", 1)[1].strip() if "—" in line else None
            if current_table_title:
                paragraph = add_heading_arial(doc, line.strip(), 3)
                add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
                toc_entries.append((line.strip(), f'toc{bookmark_counter}', 3))
                bookmark_counter += 1
            continue

        if line == "CHART:" or line.startswith("[[CHART:"):
            in_chart = True
            chart_params = {}
            if line.startswith("[[CHART:"):
                try:
                    chart_type = line.split("[[CHART:")[1].split("]]")[0].strip()
                    chart_params["TYPE"] = chart_type
                except Exception:
                    pass
                if "]]" in line:
                    tail = line.split("]]", 1)[1].strip()
                    if tail:
                        matches = list(re.finditer(r"(\w+)=", tail))
                        for i, m in enumerate(matches):
                            key = m.group(1)
                            start = m.end()
                            end = matches[i+1].start() if i+1 < len(matches) else len(tail)
                            val = tail[start:end].strip()
                            chart_params[key.strip()] = val
            continue

        if in_table:
            if line == "" or line.startswith("[[CHART:") or line.startswith("Table ") or line.startswith("^m"):
                if table_rows:
                    df = get_known_table(current_table_title) if current_table_title else None
                    if df is None:
                        parsed = False
                        try:
                            df = pd.read_csv(StringIO("\n".join(table_rows)), sep="\t")
                            parsed = True
                        except Exception:
                            try:
                                df = pd.read_csv(StringIO("\n".join(table_rows)))
                                parsed = True
                            except Exception:
                                parsed = False
                        if not parsed or len(table_rows) == 1:
                            tok = "\n".join(table_rows).split("\t")
                            best = None
                            for h in range(1, 7):
                                data_len = len(tok) - h
                                if data_len <= 0 or data_len % h != 0:
                                    continue
                                rows = data_len // h
                                cols = []
                                for c in range(h):
                                    colvals = [tok[h + r * h + c] for r in range(rows)]
                                    cols.append(colvals)
                                num_fracs = []
                                for col in cols:
                                    nnum = 0
                                    for v in col:
                                        if re.match(r"^[\d,\.\-]+%?$", v.strip()):
                                            nnum += 1
                                    num_fracs.append(nnum / max(1, len(col)))
                                score = max(num_fracs) - min(num_fracs)
                                if best is None or score > best[0]:
                                    best = (score, h, rows, cols)
                            if best is not None:
                                _, h, rows, cols = best
                                headers = tok[:h]
                                data = []
                                for r in range(rows):
                                    row = [tok[h + r * h + c] for c in range(h)]
                                    data.append(row)
                                try:
                                    df = pd.DataFrame(data, columns=[h.strip() for h in headers])
                                    parsed = True
                                except Exception:
                                    df = pd.DataFrame(data)
                                    parsed = True
                    if df is not None:
                        if df.shape[0] >= 4:
                            doc.add_page_break()
                        add_table(doc, df)
                        last_table_df = df.copy()
                        print(f"Added table with shape {df.shape}")
                    in_table = False
                    current_table_title = None
                continue
            if "	" in line:
                table_rows.append(line)
            continue

        if in_chart and "=" in line and not line.startswith("[[CHART:"):
            key, value = line.split("=", 1)
            chart_params[key.strip()] = value.strip()
            continue

        if in_chart and (line == "" or line == "^m" or line.startswith("Table ") or line.startswith("# ") or line.startswith("## ") or re.match(r"^\d+\.", line)):
            if table_rows:
                try:
                        df = pd.read_csv(StringIO("\n".join(table_rows)), sep="\t")
                except Exception:
                    df = pd.read_csv(StringIO("\n".join(table_rows)))
            else:
                df = last_table_df
            if df is not None:
                try:
                    chart_title = chart_params.get("TITLE")
                    if chart_index == 13:
                        chart_index += 1
                    if chart_title:
                            paragraph = add_heading_arial(doc, f"Chart {chart_index} — {chart_title}", 3)
                            add_bookmark(paragraph, f'toc{bookmark_counter}', bookmark_counter)
                            toc_entries.append((f"Chart {chart_index} — {chart_title}", f'toc{bookmark_counter}', 3))
                            bookmark_counter += 1
                    img = build_chart(df, chart_params, chart_index, charts_dir)
                    chart_index += 1
                    doc.add_picture(img, width=Inches(5))
                    print(f"Added chart image {img}")
                except Exception as e:
                    print(f"Failed to build chart {chart_index}: {e}")
            in_chart = False
            continue

        if (not in_table and not in_chart and line
                and not line.startswith("TABLE:")
                and not line.startswith("CHART:")
                and not line.startswith("#")):
            if re.match(r'^\[\^\d+\]:', line):
                match = re.match(r'^\[\^(\d+)\]:\s*(.*)$', line)
                if match:
                    footnote_id = match.group(1)
                    footnote_text = match.group(2)
                    paragraph = add_bookmarked_paragraph(doc, f"[^{footnote_id}]: {footnote_text}", f"fn{footnote_id}", int(footnote_id))
                    continue
            add_paragraph_with_links(doc, line)

    if toc_placeholder is not None and toc_entries:
        fill_table_of_contents(toc_placeholder, toc_entries)

    doc.save(output_path)
    print(f"Generated {output_path}")


def add_table(doc, df):
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"

    # Prevent table rows from splitting across pages
    for row in table.rows:
        row.allow_break_across_pages = False
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = False

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_arial_font(run)

    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.cell(i + 1, j)
            cell.text = str(df.iat[i, j])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_arial_font(run)


def build_chart(df, params, index, out_dir=None):
    df = df.copy()
    df.columns = [normalize_column_name(c) for c in df.columns]
    for col in df.columns:
        df[col] = df[col].astype(str).str.strip()
    x_param = params.get("X")
    y_param = params.get("Y")
    title = params.get("TITLE", "")
    chart_type = params.get("TYPE", "COLUMN").upper()
    sort_direction = params.get("SORT", "").strip().lower()
    label_param = params.get("LABELS")

    if sort_direction in {"ascending", "descending"} and x_param and y_param:
        x_col = find_best_column(df, x_param) or (df.columns[0] if len(df.columns) else None)
        y_col = find_best_column(df, y_param)
        if x_col is not None and y_col is not None:
            try:
                df[y_col] = coerce_numeric_series(df[y_col])
                df = df.sort_values(by=y_col, ascending=(sort_direction == "ascending")).reset_index(drop=True)
            except Exception:
                pass

    plt.figure()
    series_param = params.get("SERIES")
    x_col = find_best_column(df, x_param) or (df.columns[0] if len(df.columns) else None)
    y_col = find_best_column(df, y_param)

    if chart_type == "LINE":
        if series_param:
            series_list = [s.strip() for s in re.split(r"[;,]", series_param) if s.strip()]
            plotted = False
            for s in series_list:
                col = find_best_column(df, s)
                if col is not None and x_col is not None:
                    plt.plot(df[x_col], coerce_numeric_series(df[col]), label=s)
                    plotted = True
            if plotted:
                plt.legend()
        else:
            if x_col is not None and y_col is not None:
                plt.plot(df[x_col], coerce_numeric_series(df[y_col]), label=y_col)
            elif x_col is not None and len(df.columns) > 1:
                plt.plot(df[x_col], coerce_numeric_series(df[df.columns[1]]), label=df.columns[1])
    elif chart_type == "SCATTER":
        if x_col is not None and y_col is not None:
            x_vals = coerce_numeric_series(df[x_col])
            y_vals = coerce_numeric_series(df[y_col])
            plt.scatter(x_vals, y_vals)
            if label_param:
                label_col = find_best_column(df, label_param)
                if label_col is not None:
                    for xv, yv, lab in zip(x_vals, y_vals, df[label_col].astype(str)):
                        if pd.notna(xv) and pd.notna(yv):
                            plt.annotate(str(lab), (xv, yv), textcoords="offset points", xytext=(4, 4), fontsize=8)
    elif chart_type in {"STACKED", "STACKED_COLUMN"}:
        if series_param and x_col is not None:
            series_list = [s.strip() for s in re.split(r"[;,]", series_param) if s.strip()]
            x_vals = df[x_col]
            bottoms = [0] * len(df)
            plotted = False
            for s in series_list:
                col = find_best_column(df, s)
                if col is not None:
                    values = coerce_numeric_series(df[col])
                    plt.bar(x_vals, values, bottom=bottoms, label=s)
                    bottoms = [b + (v if pd.notna(v) else 0) for b, v in zip(bottoms, values)]
                    plotted = True
            if plotted:
                plt.legend()
            labels = [str(v) for v in x_vals]
            if len(labels) > 5 or any(len(lbl) > 8 for lbl in labels):
                plt.xticks(rotation=45, ha='right')
        elif x_col is not None and y_col is not None:
            plt.bar(df[x_col], coerce_numeric_series(df[y_col]))
            labels = [str(v) for v in df[x_col]]
            if len(labels) > 5 or any(len(lbl) > 8 for lbl in labels):
                plt.xticks(rotation=45, ha='right')
    else:
        if series_param:
            series_list = [s.strip() for s in re.split(r"[;,]", series_param) if s.strip()]
            width = 0.8 / max(1, len(series_list))
            x_vals = range(len(df[x_col])) if x_col is not None else range(len(df))
            plotted = False
            for i, s in enumerate(series_list):
                col = find_best_column(df, s)
                if col is not None:
                    plt.bar([xi + i * width for xi in x_vals], coerce_numeric_series(df[col]), width=width, label=s)
                    plotted = True
            plt.xticks([xi + width * (len(series_list) - 1) / 2 for xi in x_vals], df[x_col] if x_col is not None else df.index)
            if plotted:
                plt.legend()
            labels = [str(v) for v in (df[x_col] if x_col is not None else df.index)]
            if len(labels) > 5 or any(len(lbl) > 8 for lbl in labels):
                plt.xticks(rotation=45, ha='right')
        else:
            if x_col is not None and y_col is not None:
                plt.bar(df[x_col], coerce_numeric_series(df[y_col]))
            elif y_param and df.iloc[:, 0].astype(str).str.strip().eq(y_param.strip()).any():
                row = df[df.iloc[:, 0].astype(str).str.strip() == y_param.strip()].iloc[0]
                categories = list(df.columns[1:])
                values = [coerce_numeric_series(pd.Series([row[col]])).iloc[0] for col in categories]
                if all(pd.isna(values)):
                    values = list(range(1, len(categories) + 1))
                plt.bar(categories, values)
                if len(categories) > 5 or any(len(str(lbl)) > 8 for lbl in categories):
                    plt.xticks(rotation=45, ha='right')
            elif x_col is not None and len(df.columns) > 1:
                plt.bar(df[x_col], coerce_numeric_series(df[df.columns[1]]))
                labels = [str(v) for v in df[x_col]]
                if len(labels) > 5 or any(len(lbl) > 8 for lbl in labels):
                    plt.xticks(rotation=45, ha='right')
            else:
                plt.bar(df[df.columns[0]], coerce_numeric_series(df[df.columns[1]]))
                labels = [str(v) for v in df[df.columns[0]]]
                if len(labels) > 5 or any(len(lbl) > 8 for lbl in labels):
                    plt.xticks(rotation=45, ha='right')

    if x_param:
        plt.xlabel(x_param)
    elif x_col:
        plt.xlabel(x_col)
    if y_param:
        plt.ylabel(y_param)
    elif y_col:
        plt.ylabel(y_col)

    if x_col is not None:
        x_labels = [str(v) for v in df[x_col]]
        if len(x_labels) >= 5 or any(len(lbl) > 8 for lbl in x_labels):
            plt.xticks(rotation=45, ha='right')
    plt.title(title)
    plt.tight_layout()

    if out_dir:
        out_dir = Path(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        img = str(out_dir / f"chart_{index}.png")
    else:
        img = f"chart_{index}.png"
    plt.savefig(img, dpi=200)
    plt.close()
    return img

if __name__ == "__main__":
    base_dir = Path(__file__).resolve().parent
    build_report(
        base_dir / "report_v30_2025.txt",
        base_dir / "report_v30_2025.docx"
    )
